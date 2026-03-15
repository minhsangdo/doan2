from docx import Document

doc = Document('ke_hoach_do_an_dnc_chatbot.docx')

with open('docx_output.txt', 'w', encoding='utf-8') as f:
    f.write("=== PARAGRAPHS ===\n")
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip():
            f.write(f"P{i}: {para.text}\n")

    f.write("\n=== TABLES ===\n")
    for t_idx, table in enumerate(doc.tables):
        f.write(f"\n--- Table {t_idx} ---\n")
        for r_idx, row in enumerate(table.rows):
            row_data = []
            for cell in row.cells:
                txt = cell.text.strip().replace("\n", " | ")
                row_data.append(txt)
            sep = " || "
            f.write(f"  R{r_idx}: {sep.join(row_data)}\n")

print("Done - output saved to docx_output.txt")
