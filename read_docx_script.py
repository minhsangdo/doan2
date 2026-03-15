from docx import Document

doc = Document('ke_hoach_do_an_dnc_chatbot.docx')

print("=== PARAGRAPHS ===")
for i, para in enumerate(doc.paragraphs):
    if para.text.strip():
        print(f"P{i}: {para.text}")

print("\n=== TABLES ===")
for t_idx, table in enumerate(doc.tables):
    print(f"\n--- Table {t_idx} ---")
    for r_idx, row in enumerate(table.rows):
        row_data = []
        for cell in row.cells:
            txt = cell.text.strip().replace("\n", " | ")
            row_data.append(txt)
        sep = " || "
        print(f"  R{r_idx}: {sep.join(row_data)}")
